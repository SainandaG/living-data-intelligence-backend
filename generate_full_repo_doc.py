"""
generate_full_repo_doc.py
Comprehensive Architectural & Functional Documentation — Living Data Intelligence Platform
Target: ~20 pages
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ─── Colour palette ───────────────────────────────────────────────────────────
C_DARK_BLUE  = (0,   51,  102)
C_MID_BLUE   = (0,   84,  166)
C_TEAL       = (0,  128,  128)
C_DARK_GRAY  = (50,  50,  50)
C_MID_GRAY   = (100, 100, 100)
C_WHITE      = (255, 255, 255)
C_HEADER_BG  = "003366"
C_ALT_BG     = "E8F0F8"

# ─── Style helpers ────────────────────────────────────────────────────────────
def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*C_DARK_BLUE)
        run.font.name = "Calibri"
        run.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*C_MID_BLUE)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*C_TEAL)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def para(text="", indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        set_font(r, bold=True, color=C_MID_BLUE, size=10.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        set_font(r, bold=True, color=C_MID_BLUE, size=10.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EEF4FB")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5a)
    return p

def table(headers, rows, col_widths=None, alt_rows=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        _shade_cell(hc[i], C_HEADER_BG)
        for run in hc[i].paragraphs[0].runs:
            set_font(run, bold=True, size=10, color=C_WHITE)
    for ri, row_data in enumerate(rows):
        rc = t.add_row().cells
        fill = C_ALT_BG if (alt_rows and ri % 2 == 1) else "FFFFFF"
        for i, val in enumerate(row_data):
            rc[i].text = str(val)
            _shade_cell(rc[i], fill)
            for run in rc[i].paragraphs[0].runs:
                set_font(run, size=10)
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def divider(color="003366"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"),   "single")
    bt.set(qn("w:sz"),    "6")
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), color)
    pBdr.append(bt)
    pPr.append(pBdr)

def info_box(text, fill="E8F0F8"):
    """A shaded info box paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)
    r = p.add_run(text)
    set_font(r, italic=True, size=10, color=C_DARK_GRAY)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("LIVING DATA INTELLIGENCE PLATFORM")
set_font(r, size=26, bold=True, color=C_DARK_BLUE)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Comprehensive Architectural & Functional Documentation")
set_font(r, size=15, color=C_MID_BLUE)

doc.add_paragraph()
vp = doc.add_paragraph()
vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = vp.add_run("Complete Repository Reference  ·  Version 1.0")
set_font(r, size=11, italic=True, color=C_MID_GRAY)

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dp.add_run("March 2026")
set_font(r, size=11, italic=True, color=C_MID_GRAY)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run(
    "Transform database schemas into interactive 3D visualisations\n"
    "with real-time transaction monitoring, autonomous AI agents,\n"
    "machine-learning analytics, and voice control."
)
set_font(r, size=11, color=C_DARK_GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc = [
    ("1.",   "Executive Summary & Platform Vision"),
    ("2.",   "System Architecture Overview"),
    ("3.",   "Technology Stack"),
    ("4.",   "Backend Application — FastAPI Entry Point"),
    ("  4.1","Router Registry & API Surface"),
    ("  4.2","Middleware Stack"),
    ("  4.3","Background Tasks & Lifecycle"),
    ("5.",   "Database Connectivity Layer"),
    ("  5.1","Multi-Dialect Connection Management"),
    ("  5.2","Identifier Validation & SQL Injection Guard"),
    ("  5.3","Neon Serverless & Keep-Alive"),
    ("6.",   "Neural Core — Active Schema Intelligence"),
    ("  6.1","Gravity & Vitality Metrics"),
    ("  6.2","WEZU Energy Domain Ontology"),
    ("  6.3","Sub-modules: Scanner, Metrics, Signal Processor, Reporter"),
    ("7.",   "Autonomous Agent Architecture"),
    ("  7.1","T0 Agent — Voice Intent Brain"),
    ("  7.2","T1 Agent — Action Execution Engine"),
    ("  7.3","WEZU Domain Agents"),
    ("  7.4","Agent Service & Autonomous Loop"),
    ("8.",   "Graph Generation Engine"),
    ("  8.1","Statistical 3D Positioning"),
    ("  8.2","Cluster Strategies: Heuristic vs. NetworkX"),
    ("  8.3","Node & Edge Rendering"),
    ("9.",   "Real-Time WebSocket Protocol"),
    ("  9.1","Server → Client Message Types"),
    ("  9.2","Client → Server Message Types"),
    ("  9.3","Heartbeat, Presence & Multiplayer"),
    ("10.",  "Intelligence Services Layer"),
    ("  10.1","Schema Analyser"),
    ("  10.2","Anomaly Detector"),
    ("  10.3","Pattern Analyser"),
    ("  10.4","Predictive Engine"),
    ("  10.5","Root Cause Analyser"),
    ("  10.6","Causal Intelligence"),
    ("  10.7","Recommendation Engine"),
    ("11.",  "ML Analytics — Work on Data"),
    ("  11.1","Four Algorithm Families"),
    ("  11.2","Intelligence Pipeline"),
    ("  11.3","API Endpoints"),
    ("12.",  "Frontend Architecture"),
    ("  12.1","Component Tree"),
    ("  12.2","3D Visualisation Subsystem (ThreeGraph)"),
    ("  12.3","Latent Space Overlay"),
    ("  12.4","State Management (Zustand Stores)"),
    ("  12.5","Audio Engine"),
    ("  12.6","Voice Control"),
    ("13.",  "Authentication & Security"),
    ("14.",  "Configuration & Deployment"),
    ("15.",  "Test Suite"),
    ("16.",  "Complete API Route Catalogue"),
    ("17.",  "Future Roadmap"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{num}  {title}")
    is_top = not num.startswith("  ")
    set_font(r, size=10.5 if is_top else 10,
             bold=is_top,
             color=C_DARK_BLUE if is_top else C_DARK_GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary & Platform Vision")
divider()
para(
    "The Living Data Intelligence Platform transforms static relational database schemas "
    "into living, breathing 3D environments. It connects to any PostgreSQL, MySQL, or MongoDB "
    "database and immediately begins a continuous cycle of introspection, analysis, and "
    "visualisation — without writing a single byte back to the source database."
)
para(
    "The platform is organised around five pillars:"
)
bullet("Spatial Intelligence — A Three.js / React-Three-Fiber WebGL engine renders every table as a node in a 3D gravitational graph, with edges representing foreign-key relationships.", bold_prefix="01 Spatial Intelligence")
bullet("Autonomous Agents — A T0 voice-to-intent brain and a T1 action execution engine work in tandem to let users control the graph with natural language or voice commands.", bold_prefix="02 Autonomous Agents")
bullet("Neural Core — An always-on Active Schema Intelligence service continuously scans schema metadata, calculates gravity/vitality metrics, and feeds the visualisation.", bold_prefix="03 Neural Core")
bullet("Real-Time Telemetry — A WebSocket bus broadcasts live TPS, evolved node states, and anomaly alerts every 2 seconds to every connected browser tab.", bold_prefix="04 Real-Time Telemetry")
bullet("ML Analytics (Work on Data) — A read-only scikit-learn pipeline runs Classification, Regression, Time Series, and Clustering directly on database table rows.", bold_prefix="05 ML Analytics")

para(
    "The platform is built on a strict Zero-Write Policy: all SQL is SELECT-only, "
    "all computation happens in server RAM, and no temporary tables are ever created "
    "in the user's database."
)
info_box(
    "Primary target audiences: database architects, data engineers, data scientists, "
    "and enterprise observability teams who need an intelligent, visual interface to "
    "their production data assets."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Architecture Overview")
divider()
para("The system is a three-tier web application with a dedicated AI/LLM layer:")

arch_layers = [
    ("Tier 1 — Presentation",
     "React 18 SPA (Vite + TypeScript/JSX), Three.js/R3F 3D engine, Recharts 2D analytics, Zustand state management. Served from http://localhost:5173 in development."),
    ("Tier 2 — Application",
     "FastAPI (Python 3.8+) async server on Uvicorn. Exposes REST + WebSocket APIs. Hosts Neural Core, Agent system, ML pipeline, and all intelligence services. Served on port 8001."),
    ("Tier 3 — Data",
     "Read-only connection pools to PostgreSQL (asyncpg), MySQL (aiomysql), and MongoDB (pymongo). The platform never modifies data."),
    ("AI/LLM Layer",
     "Google Gemini API (GOOGLE_API_KEY) for natural-language chat, intent classification, and schema explanation. Low-temperature (0.2) for deterministic analytical output."),
]
table(
    ["Tier", "Description"],
    arch_layers,
    col_widths=[Inches(1.8), Inches(5.0)],
)

h2("2.1  High-Level Data Flow")
steps = [
    "Browser → POST /api/database/connect → DatabaseConnector creates async pool.",
    "POST /api/schema/analyze → SchemaAnalyser introspects tables, columns, FK relationships.",
    "POST /api/graph/generate → GraphGenerator builds node-edge JSON with 3D coordinates.",
    "GET /api/graph/{id} → Frontend renders Three.js scene; nodes pulse by vitality.",
    "WS /ws/{connection_id} → RealtimeMonitor streams metrics_update every 2 seconds.",
    "Voice input → T0Agent classifies intent → T1Agent executes graph action.",
    "POST /api/ml/analyze → ML pipeline trains scikit-learn model → returns insights.",
]
for s in steps:
    numbered(s)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Technology Stack")
divider()

h2("3.1  Backend")
backend_stack = [
    ("FastAPI",           "2.x",   "Async Python web framework; automatic OpenAPI/Swagger docs."),
    ("Uvicorn",           "latest","ASGI server; hot-reload in development mode."),
    ("asyncpg",           "latest","Async PostgreSQL driver; connection pool management."),
    ("aiomysql",          "latest","Async MySQL driver."),
    ("pymongo",           "latest","MongoDB driver (sync, wrapped in executor)."),
    ("scikit-learn",      "1.x",   "ML algorithms: Random Forest, SVM, KNN, Logistic Regression, Ridge, Lasso, GradientBoosting, KMeans, DBSCAN."),
    ("NumPy / pandas",    "latest","Numerical computation; DataFrame preprocessing."),
    ("NetworkX",          "latest","Graph theory: Louvain community detection, PageRank."),
    ("python-jose",       "latest","JWT token creation and verification."),
    ("passlib + bcrypt",  "latest","Password hashing."),
    ("slowapi",           "latest","Rate limiting middleware."),
    ("python-dotenv",     "latest","Environment variable loading."),
    ("aiofiles",          "latest","Async file I/O for connection debug logging."),
]
table(
    ["Library", "Version", "Purpose"],
    backend_stack,
    col_widths=[Inches(1.5), Inches(0.8), Inches(4.5)],
)

h2("3.2  Frontend")
frontend_stack = [
    ("React 18",              "18.x",  "Component model; hooks-based state; Concurrent Mode."),
    ("Vite",                  "5.x",   "Build tool; fast HMR; ESM-native bundling."),
    ("Three.js / R3F",        "latest","3D WebGL scene; React-Three-Fiber declarative wrapper."),
    ("@react-three/drei",     "latest","Camera, orbit controls, helpers for R3F."),
    ("d3-force-3d",           "latest","Physics simulation for force-directed graph layout."),
    ("Recharts",              "2.x",   "2D chart library for scatter, bar, line, heatmap plots."),
    ("Zustand",               "4.x",   "Lightweight global state management (6 stores)."),
    ("framer-motion",         "latest","Animation library; smooth transitions in modals."),
    ("Lucide React",          "latest","Icon library used throughout the UI."),
    ("TypeScript",            "5.x",   "Type-safe agents, hooks, and utilities."),
    ("Web Audio API",         "native","Procedural sound engine for graph events."),
    ("Web Speech API",        "native","Voice recognition for voice control feature."),
]
table(
    ["Library", "Version", "Purpose"],
    frontend_stack,
    col_widths=[Inches(1.9), Inches(0.8), Inches(4.1)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. BACKEND APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Backend Application — FastAPI Entry Point")
divider()
para(
    "The entry point backend/main.py bootstraps the entire platform. "
    "It uses FastAPI's asynccontextmanager lifespan hook to manage startup and "
    "graceful shutdown of all background tasks and database connections."
)

h2("4.1  Router Registry & API Surface")
para(
    "All route registrations are extracted to backend/router_registry.py via "
    "register_all_routes(app, registry, expected_routers). Routers are classified as "
    "Required (app will not start if missing) or Optional (app starts in degraded mode)."
)

required_routes = [
    ("app.api.auth",         "/api/auth",     "Login, refresh-token, logout endpoints."),
    ("app.api.database",     "/api",          "Connect, disconnect, list connections."),
    ("app.api.schema",       "/api",          "Schema analysis, table listing, column metadata."),
    ("app.api.graph",        "/api",          "Graph generation, optimisation, cluster update."),
    ("app.api.metrics",      "/api",          "Real-time metrics endpoint."),
    ("app.api.drilldown",    "/api",          "Drill-down into table record level."),
    ("app.api.hierarchy",    "/api",          "Hierarchical flow analysis."),
    ("app.api.internal_node","/api",          "Internal node X-Ray deep inspection."),
    ("app.api.ai",           "/api/ai",       "AI chat, explain, classify endpoints."),
    ("app.api.agent",        "(root)",        "T0/T1 agent voice command endpoint."),
    ("app.api.websocket",    "/ws",           "WebSocket stream endpoint."),
]
table(
    ["Module", "Prefix", "Purpose"],
    required_routes,
    col_widths=[Inches(2.0), Inches(1.0), Inches(3.8)],
)

h3("Optional Routers")
optional_routes = [
    ("app.api.latent_stream",   "Latent space streaming SSE."),
    ("app.api.data_explorer",   "Paginated table data browsing."),
    ("app.api.data_flow",       "Data flow analysis between tables."),
    ("app.api.chat",            "Persistent AI analyst chat sessions."),
    ("app.api.evolution",       "Schema evolution timeline playback."),
    ("app.api.ml",              "Legacy ML endpoints."),
    ("app.api.ml_analysis",     "Work on Data — scikit-learn ML analysis."),
    ("app.api.events",          "Transaction event processing."),
    ("app.api.explainability",  "XAI explanation generation."),
    ("app.api.vitals",          "System vitals dashboard endpoint."),
    ("app.api.intelligence",    "Full intelligence hub: patterns, predictions, anomalies."),
    ("app.api.ontology",        "Domain ontology explorer."),
    ("app.api.node_xray",       "Node X-Ray panel endpoint."),
    ("app.api.simulation",      "Data simulator seeding endpoint."),
    ("app.api.seeder_api",      "Database seed data generation."),
]
table(
    ["Module", "Purpose"],
    optional_routes,
    col_widths=[Inches(2.2), Inches(4.6)],
)

h2("4.2  Middleware Stack")
bullet("CORSMiddleware — reads ALLOWED_ORIGINS / CORS_ORIGINS from .env; supports JSON-list and comma-separated formats; auto-adds localhost:5173 in development.")
bullet("GZipMiddleware — compresses responses ≥ 1,000 bytes; reduces bandwidth on large graph payloads.")
bullet("Request Logging Middleware — injects a random 8-char request_id (UUID) per request; logs method, path, status, and duration_ms.")
bullet("SlowAPI Rate Limiter — IP-based rate limiting on auth endpoints via slowapi.")
bullet("Global Exception Handler — catches all unhandled exceptions; returns structured JSON {error, code, path} without leaking stack traces.")

h2("4.3  Background Tasks & Lifecycle")
bg_tasks = [
    ("stream_metrics()",                 "WebSocket background broadcaster — pushes metrics to all connected clients every 2 seconds."),
    ("keep_alive_task()",                "Sends SELECT 1 to every active connection every 4 minutes to prevent serverless cold-start drops."),
    ("agent_service.start_autonomous_loop()", "Runs the autonomous agent monitoring loop — continuously checks schema for changes and anomalies."),
]
table(
    ["Task", "Description"],
    bg_tasks,
    col_widths=[Inches(2.5), Inches(4.3)],
)
para("On shutdown, all tasks are cancelled via asyncio.gather with a 3-second timeout, then all DB pools are drained with a 5-second timeout.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. DATABASE CONNECTIVITY LAYER
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Database Connectivity Layer")
divider()
para("File: backend/app/services/db_connector.py — DatabaseConnector singleton.")

h2("5.1  Multi-Dialect Connection Management")
dialects = [
    ("PostgreSQL",  "asyncpg",  "Async connection pool; native PostgreSQL protocol; Neon SSL override."),
    ("MySQL",       "aiomysql", "Async connection pool; charset=utf8mb4 by default."),
    ("MongoDB",     "pymongo",  "Sync MongoClient wrapped in ThreadPoolExecutor for async compatibility."),
    ("Neon",        "asyncpg",  "PostgreSQL over TLS (sslmode=require); auto-detected by hostname containing 'neon.tech'."),
]
table(
    ["Dialect", "Driver", "Notes"],
    dialects,
    col_widths=[Inches(1.0), Inches(1.0), Inches(4.8)],
)
para("Each connection is assigned a sequential ID (conn_1, conn_2, …). The connection dictionary is stored in self.connections keyed by connection_id.")

h2("5.2  Identifier Validation & SQL Injection Guard")
para(
    "validate_identifier(name) rejects any identifier that does not match "
    r"the regex ^[a-zA-Z0-9_\.]+$, raising ValueError on injection attempts. "
    "quote_identifier(connection_id, name) wraps validated identifiers in "
    "backtick quotes (MySQL) or double-quote quotes (PostgreSQL/others)."
)
code(
    "# MySQL:      `table_name`\n"
    "# PostgreSQL: \"table_name\"\n"
    "# All user-supplied identifiers MUST pass through this function."
)

h2("5.3  Neon Serverless & Keep-Alive")
para(
    "Neon Postgres databases enter a sleep state after inactivity. "
    "The connector handles the following Neon-specific wake-up errors with "
    "exponential-backoff retry (up to NEON_MAX_RETRIES=3, NEON_CONNECT_TIMEOUT=180s):"
)
bullet("'Connection refused' / 'could not connect to server'")
bullet("'the database system is starting up'")
bullet("'endpoint is disabled' / 'pg_sleep'")
para("The keep_alive_task() in main.py sends SELECT 1 every 4 minutes to prevent cold starts in the first place.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. NEURAL CORE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Neural Core — Active Schema Intelligence")
divider()
para(
    "File: backend/app/services/neural_core/core.py — NeuralCore singleton. "
    "Treats the database schema not as a static structure but as a living organism "
    "that emits signals. It is the primary data source for graph positioning and animation."
)

h2("6.1  Gravity & Vitality Metrics")
metrics = [
    ("Gravity (G)",  "G = log₁₀(row_count + 1) × centrality_weight",
     "Importance score. High-gravity nodes are positioned centrally in the 3D graph and rendered with larger radii."),
    ("Vitality (V)", "Influenced by recent transaction rates and anomaly detections.",
     "Health/activity level. Node 'pulse' animation frequency in the UI is tied to this value."),
    ("Entropy (E)",  "Schema complexity and relationship fragmentation.",
     "Measure of how interconnected and complex the table's role is."),
    ("Hub Score",    "Computed via graph centrality (in-degree + out-degree weighting).",
     "Identifies tables that act as data hubs — many other tables reference them."),
]
table(
    ["Metric", "Formula / Logic", "Visualisation Effect"],
    metrics,
    col_widths=[Inches(0.9), Inches(2.5), Inches(3.4)],
)

h2("6.2  WEZU Energy Domain Ontology")
para(
    "The NeuralCore ships with a hard-coded WEZU_ENERGY_ONTOLOGY dictionary that "
    "provides domain-aware intelligence for electric-vehicle / battery-swap deployments. "
    "Tables matching these names receive elevated gravity weights and specific type labels:"
)
wezu_tables = [
    ("batteries",         10.0, "asset",          "Primary storage unit — critical for energy distribution."),
    ("stations",           9.0, "infrastructure", "Physical swap points — core distribution network nodes."),
    ("iot_devices",        8.0, "asset",          "Remote sensors — real-time monitoring & predictive maintenance."),
    ("telematics_data",    7.0, "telemetry",      "High-velocity operational metrics — anomaly detection basis."),
    ("battery_health_log", 8.0, "telemetry",      "SOH historical ledger — lifecycle analysis."),
    ("gps_tracking_log",   7.0, "telemetry",      "Spatial movement data — geographical pattern analysis."),
    ("swap_transactions",  7.0, "transaction",    "Asset exchange event ledger — network utilisation metric."),
    ("wallet_transactions",6.0, "financial",      "Monetary flow — reconciliation and fraud detection."),
]
table(
    ["Table Pattern", "Gravity", "Type", "Justification"],
    wezu_tables,
    col_widths=[Inches(1.6), Inches(0.6), Inches(1.1), Inches(3.5)],
)

h2("6.3  Sub-modules")
sub_modules = [
    ("schema_scanner.py",    "Introspects live schema: tables, columns, FK relationships, row counts."),
    ("metrics_calculator.py","Computes gravity, hub score, in/out degree, entropy per table."),
    ("signal_processor.py",  "Processes incremental schema change signals and updates snapshots."),
    ("analysis_reporter.py", "Formats analysis results into structured JSON for API responses."),
]
table(
    ["File", "Responsibility"],
    sub_modules,
    col_widths=[Inches(2.0), Inches(4.8)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. AUTONOMOUS AGENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Autonomous Agent Architecture")
divider()
para(
    "The platform employs a hierarchical multi-agent system with dedicated "
    "TypeScript agents on the frontend and Python agents on the backend, "
    "communicating via the same WebSocket bus."
)

agent_tier = [
    ("T0", "Voice Intent Brain",   "Process voice transcription, classify intent, maintain context, dispatch to T1."),
    ("T1", "Action Execution Engine", "Execute graph/analytics/UI actions, manage state transitions, report results."),
    ("WEZU", "Domain-Specific Intelligence", "Battery degradation monitoring (SoH), GPS geofence breach detection, energy-specific anomaly patterns."),
]
table(
    ["Tier", "Role", "Responsibilities"],
    agent_tier,
    col_widths=[Inches(0.5), Inches(1.8), Inches(4.5)],
)

h2("7.1  T0 Agent — Voice Intent Brain")
para("File: backend/app/agents/t0_agent.py  (frontend mirror: frontend/src/agents/T0Agent.ts)")
para("T0 processes voice transcription text through a three-phase pipeline:")
numbered("Intent Classification via UnifiedIntentClassifier — uses regex + LLM to map natural language to a structured intent object.")
numbered("Context Maintenance — rolling window of last 5 utterances (max_context=5) stored in ContextManager.")
numbered("Dispatch — sends classified intent to T1 for action execution.")
para("Feature flag USE_ENHANCED_T0_AGENT enables V2 features (richer context, multi-turn conversation).")

h2("7.2  T1 Agent — Action Execution Engine")
para("File: backend/app/agents/t1_agent.py  (frontend mirror: frontend/src/agents/T1Agent.ts)")
t1_actions = [
    ("graph.highlight",           "Highlight a specific node in the 3D graph."),
    ("graph.zoom_cluster",        "Fly the camera to a specific cluster."),
    ("graph.start_flow",          "Begin animated particle flow on edges."),
    ("graph.stop_flow",           "Stop particle flow animation."),
    ("graph.recalculate_gravity", "Trigger gravity recalculation for all nodes."),
    ("graph.reset_view",          "Return camera to default position."),
    ("analytics.anomaly",         "Run anomaly detection on the active schema."),
    ("analytics.cluster",         "Apply or re-apply clustering to the graph."),
    ("ui.show_schema",            "Open the schema panel for a table."),
    ("ui.drill_down",             "Enter drill-down mode for a table's records."),
    ("graph.start_evolution",     "Begin schema evolution timeline playback."),
    ("graph.stop_evolution",      "Stop evolution playback."),
    ("graph.simulate_formation",  "Run node formation simulation animation."),
]
table(
    ["Action ID", "Effect"],
    t1_actions,
    col_widths=[Inches(2.2), Inches(4.6)],
)

h2("7.3  WEZU Domain Agents")
para("File: backend/app/services/wezu_agents.py")
para(
    "WEZU agents are specialised intelligence modules activated when the schema "
    "matches the WEZU Energy ontology. They perform:"
)
bullet("State-of-Health (SoH) monitoring — detects battery degradation patterns by analysing battery_health_log trends.")
bullet("GPS Geofence Breach detection — identifies anomalous location patterns in gps_tracking_log against known station boundaries.")
bullet("Swap Anomaly detection — flags unusually high or low battery swap rates at stations.")

h2("7.4  Agent Service & Autonomous Loop")
para("File: backend/app/agents/agent_service.py")
para(
    "The AgentService runs start_autonomous_loop() as a background asyncio task. "
    "It continuously monitors all active connections for schema changes, "
    "triggers Neural Core rescans, and emits evolved_nodes and anomalies "
    "events to the WebSocket bus."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. GRAPH GENERATION ENGINE
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Graph Generation Engine")
divider()
para("File: backend/app/services/graph_generator.py — GraphGenerator class.")

h2("8.1  Statistical 3D Positioning")
para("Every table node is assigned a deterministic 3D position (x, y, z) based on three statistical axes:")
axes = [
    ("X — Data Volume",       "log₁₀(row_count + 1), normalised to [−600, +600]. Tables with more rows move outward on X."),
    ("Y — Structural Complexity", "columns + foreign_keys, normalised to [−400, +400]. Highly connected tables move up on Y."),
    ("Z — Neural Importance", "Neural Core gravity score, normalised to [−500, +500]. Critical tables float toward +Z."),
]
table(
    ["Axis", "Formula & Effect"],
    axes,
    col_widths=[Inches(1.8), Inches(5.0)],
)

h2("8.2  Cluster Strategies")
strategies = [
    ("Heuristic (default)", "Prefix-based pattern matching (e.g., auth_ prefix groups auth_user, auth_group). Speed: instant. Accuracy: 60–80% depending on naming conventions. Uses 3-colour palette: Cyan, Green, Gold.", "Databases with consistent table naming conventions."),
    ("NetworkX (advanced)", "Louvain community detection on FK relationship graph, weighted by PageRank scores. Speed: <100ms for typical schemas. Accuracy: ~95%. Uses 4-colour palette: Blue, Purple, Orange, Green.", "Any schema; especially complex or legacy databases."),
]
table(
    ["Method", "Algorithm & Characteristics", "Best For"],
    strategies,
    col_widths=[Inches(1.2), Inches(3.6), Inches(2.0)],
)

h2("8.3  Node & Edge Rendering (Frontend)")
para("Frontend: frontend/src/components/Dashboard/ThreeGraph/")
sub_files = [
    ("SceneSetup.js",        "Initialises Three.js renderer, camera, lighting, and fog."),
    ("NodeRenderer.js",      "Creates sphere meshes per table node; applies glow shader (nodeGlowShader.glsl); animates pulse by vitality."),
    ("EdgeRenderer.js",      "Draws FK relationship edges; applies edgeFlowShader.glsl for animated particle flow."),
    ("PhysicsEngine.js",     "Runs d3-force-3d simulation for force-directed layout (Galaxy View)."),
    ("ClusterManager.js",    "Groups nodes by cluster; renders cluster boundary spheres."),
    ("InteractionHandler.js","Handles click, hover, orbit, zoom, and drag interactions."),
]
table(
    ["File", "Responsibility"],
    sub_files,
    col_widths=[Inches(1.9), Inches(4.9)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. REAL-TIME WEBSOCKET PROTOCOL
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Real-Time WebSocket Protocol")
divider()
para("File: backend/app/api/websocket.py — endpoint: WS /ws/{connection_id}")
para(
    "The WebSocket layer maintains a registry of active_connections (Dict[connection_id → List[WebSocket]]) "
    "supporting multiple browser tabs per database connection. "
    "In production mode, connections require a valid JWT token as a query parameter."
)

h2("9.1  Server → Client Message Types")
s2c = [
    ("ping",            "Heartbeat sent every 30 seconds. Payload: {type, timestamp}."),
    ("connected",       "Sent on successful handshake. Payload: {type, connection_id, client_count}."),
    ("metrics_update",  "Core telemetry burst every 2s. Contains TPS, total_rows, evolved_nodes, active_connections."),
    ("db_reconnecting", "Emitted when a Neon cold-start reconnection attempt begins."),
    ("presence_update", "Multiplayer cursor position update relayed from another client."),
    ("error",           "Protocol or database error. Payload: {type, message, code}."),
]
table(
    ["Type", "Description"],
    s2c,
    col_widths=[Inches(1.6), Inches(5.2)],
)

h2("9.2  Client → Server Message Types")
c2s = [
    ("pong",            "Heartbeat acknowledgment."),
    ("presence_update", "User sends cursor position {user_id, cursor: {x, y}} for multiplayer mode."),
    ("ping (legacy)",   "Deprecated plain-text health-check. Still handled for backward compatibility."),
]
table(
    ["Type", "Description"],
    c2s,
    col_widths=[Inches(1.6), Inches(5.2)],
)

h2("9.3  Heartbeat, Presence & Multiplayer")
bullet("safe_send(ws, payload) wraps all sends in asyncio.wait_for(timeout=5.0) — stale connections are silently dropped.")
bullet("Multi-tab support: a single database connection_id can have multiple WebSocket entries; all receive the same broadcast.")
bullet("Multiplayer: presence_update messages are relayed to all other tabs sharing the same connection_id.")
bullet("Frontend hook: frontend/src/hooks/useMultiplayer.js manages RemoteCursors.jsx overlay.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. INTELLIGENCE SERVICES LAYER
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Intelligence Services Layer")
divider()
para(
    "The platform ships a full suite of analytical intelligence services under "
    "backend/app/services/. Each service is a singleton and is exposed through "
    "the /api/intelligence router."
)

intel_services = [
    ("schema_analyzer.py",          "SchemaAnalyser",          "Primary schema introspection service. Populates NeuralCore snapshots. Caches results per connection_id."),
    ("anomaly_detector.py",         "AnomalyDetector",         "Identifies statistical outliers in row distributions and column value patterns. Feeds T1 analytics.anomaly actions."),
    ("pattern_analyzer.py",         "PatternAnalyser",         "Mines recurring data patterns (temporal, structural) across tables. Outputs pattern cards for IntelligenceHub."),
    ("predictive_engine.py",        "PredictiveEngine",        "Forecasts table growth rates and schema evolution trends using linear extrapolation."),
    ("root_cause_analyzer.py",      "RootCauseAnalyser",       "Traces anomaly origins through FK chains to identify the upstream source table."),
    ("causal_intelligence.py",      "CausalIntelligence",      "Builds causal graphs between table events using co-occurrence and temporal lag analysis."),
    ("recommendation_engine.py",    "RecommendationEngine",    "Suggests indexing strategies, schema optimisations, and analytical next-steps."),
    ("data_quality_engine.py",      "DataQualityEngine",       "Evaluates null rates, type consistency, and uniqueness ratios per column."),
    ("intelligence_engine.py",      "IntelligenceEngine",      "Orchestrator: runs all intelligence services in parallel and aggregates results for the IntelligenceHub."),
    ("xai_service.py",              "XAIService",              "Explainable AI: generates natural-language explanations for model decisions and anomaly detections."),
    ("latent_space_service.py",     "LatentSpaceService",      "Projects tables into 2D/3D semantic manifold using feature-vector similarity."),
    ("data_intelligence_analyzer.py","DataIntelligenceAnalyser","Cross-table data flow and dependency analysis."),
    ("temporal_analyzer.py",        "TemporalAnalyser",        "Time-series analysis of table growth and transaction frequencies."),
    ("data_flow_analyzer.py",       "DataFlowAnalyser",        "Traces data lineage — how records propagate between tables."),
]
table(
    ["File", "Class", "Purpose"],
    intel_services,
    col_widths=[Inches(2.2), Inches(1.6), Inches(3.0)],
)

para("Frontend dashboards consuming these services:")
dash_map = [
    ("AnomalyDashboard.jsx",      "Real-time anomaly cards with severity ratings."),
    ("PatternDashboard.jsx",      "Mined pattern cards with confidence scores."),
    ("PredictionDashboard.jsx",   "Growth forecasts with trend arrows."),
    ("RootCauseDashboard.jsx",    "FK-chain root cause trace diagrams."),
    ("RecommendationDashboard.jsx","Schema optimisation suggestion cards."),
    ("HealthDashboard.jsx",       "Overall system health score."),
    ("OntologyExplorer.jsx",      "Interactive WEZU ontology graph."),
    ("SemanticSearchDiscovery.jsx","Natural-language table/column search."),
    ("IntelligenceHub.jsx",       "Master dashboard aggregating all panels."),
]
table(
    ["Component", "Content"],
    dash_map,
    col_widths=[Inches(2.4), Inches(4.4)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. ML ANALYTICS — WORK ON DATA
# ══════════════════════════════════════════════════════════════════════════════
h1("11. ML Analytics — Work on Data")
divider()
para("File: backend/app/api/ml_analysis.py — FastAPI router prefix: /api/ml")

h2("11.1  Four Algorithm Families")
families = [
    ("Classification", "Discrete label prediction",
     "Random Forest (rf_clf), SVM, K-Nearest Neighbors (knn), Logistic Regression (logreg)",
     "Accuracy, F1, Precision, Recall"),
    ("Regression",     "Continuous value prediction",
     "Gradient Boosting (xgboost), Linear Regression, Ridge, Lasso",
     "R², RMSE, MAE"),
    ("Time Series",    "Temporal forecasting",
     "Trend + Seasonal Harmonic Regression (numpy; no statsmodels required)",
     "MAPE, RMSE, MAE, monthly_growth, trend"),
    ("Clustering",     "Unsupervised segmentation",
     "K-Means (Auto-K via Silhouette), DBSCAN (Auto-eps via k-distance 90th percentile)",
     "Silhouette Score, inertia, n_clusters, n_noise_points"),
]
table(
    ["Family", "Objective", "Algorithms", "Key Metrics"],
    families,
    col_widths=[Inches(1.1), Inches(1.4), Inches(2.4), Inches(1.9)],
)

h2("11.2  Intelligence Pipeline")
pipeline = [
    ("Stage 1: Secure Ingestion",
     "Safe-quoted SELECT query (backtick/double-quote per dialect). Row cap: min(100–5000). Column cap: 20. No DDL/DML ever generated."),
    ("Stage 2: Preprocessing",
     "Median imputation (numeric), '__missing__' fill (categorical), LabelEncoder, StandardScaler (SVM/KNN/Ridge/Lasso), np.nan_to_num sanitisation."),
    ("Stage 3: Training",
     "asyncio run_in_executor offloads scikit-learn CPU work to ThreadPoolExecutor. Prevents blocking the FastAPI event loop."),
    ("Stage 4: Insight Generation",
     "_build_insights() converts numeric metrics to natural-language strings. Family-specific condition/insight rules (F1 thresholds, MAPE bands, silhouette quality)."),
]
table(
    ["Stage", "Description"],
    pipeline,
    col_widths=[Inches(2.0), Inches(4.8)],
)

h2("11.3  API Endpoints")
ml_api = [
    ("POST /api/ml/analyze", "Run ML analysis on a table. Body: {connection_id, table, family, algo, target, features}. Returns metrics, feature_importances, predictions, insights, scatter_sample."),
    ("POST /api/ml/suggest", "Heuristic algorithm recommender. Query params: connection_id, table. Returns {suggestion, reason, table_info} with confidence score."),
]
table(
    ["Endpoint", "Description"],
    ml_api,
    col_widths=[Inches(2.0), Inches(4.8)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. FRONTEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Frontend Architecture")
divider()
para("Entry: frontend/src/main.jsx → App.jsx — React 18 SPA served by Vite on port 5173.")

h2("12.1  Component Tree")
component_groups = [
    ("Layout",          "DashboardLayout, NavigationBar, Sidebars — top-level scaffolding."),
    ("Auth",            "LoginPage — JWT login form."),
    ("Dashboard",       "ThreeGraph (3D engine), WorkOnDataModal, DeepAnalysisPage, ChatInterface, DrillDownView, DataFlowView, SchemaView, AnalyticsView, NodeXRayPanel, SystemVitalsDashboard, LatentSpaceOverlay, SemanticDiscoveryPanel, PerspectiveLineageView, LineageInsightHUD."),
    ("Intelligence",    "IntelligenceHub, AnomalyDashboard, PatternDashboard, PredictionDashboard, RootCauseDashboard, RecommendationDashboard, HealthDashboard, BlueprintOverlay, OntologyExplorer, DeepStatusDashboard."),
    ("Evolution",       "EvolutionOverlay, EvolutionMathOverlay, TimelinePlayer, NodeFormationSimulation."),
    ("WindowManager",   "ConnectionModal, Taskbar, Window — OS-like floating window system."),
    ("Panels",          "TimeMachinePanel — schema snapshot navigation."),
    ("Incident",        "WarRoomHUD — incident response command centre."),
    ("Voice",           "VoiceControl, AgentStatusPanel — Web Speech API integration."),
    ("Multiplayer",     "RemoteCursors — real-time cursor presence."),
    ("Apps",            "AnalystChat (persistent AI session), Settings."),
]
table(
    ["Group", "Components"],
    component_groups,
    col_widths=[Inches(1.2), Inches(5.6)],
)

h2("12.2  3D Visualisation Subsystem (ThreeGraph)")
para(
    "The ThreeGraph component is the visual centrepiece. It supports three visualisation modes:"
)
vis_modes = [
    ("Galaxy View",    "d3-force-3d physics simulation. Nodes repel/attract based on FK relationships and gravity scores. Best for exploring large schemas."),
    ("Latent Space",   "Tables projected into 3D semantic manifold. Clustering via heuristic (prefix) or NetworkX (Louvain). Best for discovering hidden schema groups."),
    ("Tier 3 / Voxel", "Drill-down mode: individual records rendered as 3D voxel blocks within a table's gravity well. Best for record-level inspection."),
]
table(
    ["Mode", "Description"],
    vis_modes,
    col_widths=[Inches(1.4), Inches(5.4)],
)
para("Custom GLSL shaders: nodeGlowShader.glsl (pulsing glow per node vitality), edgeFlowShader.glsl (animated particles on FK edges).")

h2("12.3  Latent Space Overlay")
para("Files: frontend/src/components/Dashboard/LatentSpace/")
bullet("LayoutEngine.js — computes 2D/3D cluster positions using t-SNE-like projection of column-type similarity vectors.")
bullet("LatentWorld.jsx — Three.js scene for the projected table space.")
bullet("LatentSpaceOverlay.jsx — React overlay with cluster labels and semantic discovery panel.")
bullet("computations.js — cosine similarity calculations for column-name vectors.")

h2("12.4  State Management (Zustand Stores)")
stores = [
    ("graphStore.js",        "Nodes, edges, cluster assignments, selected node, graph metadata."),
    ("connectionStore.js",   "Active connection_id, connection history, connection status."),
    ("realtimeStore.js",     "Live TPS, metrics_update payload, WebSocket connection state."),
    ("intelligenceStore.js", "Anomalies, patterns, predictions, recommendations from AI services."),
    ("evolutionStore.js",    "Evolution timeline snapshots and playback position."),
    ("viewStore.js",         "Active view mode (Galaxy/Latent/Voxel), camera position, sidebar state."),
    ("authStore.js",         "JWT access token, refresh token, user profile."),
]
table(
    ["Store", "State Managed"],
    stores,
    col_widths=[Inches(1.8), Inches(5.0)],
)

h2("12.5  Audio Engine")
para("Files: frontend/src/audio/ and frontend/src/utils/")
audio_files = [
    ("GraphSoundEngine.ts",         "Master audio controller — maps graph events to sound triggers."),
    ("FrequencyMapper.ts",          "Maps node metrics (gravity, vitality) to audio frequencies."),
    ("RhythmGenerator.ts",          "Generates rhythmic patterns based on TPS and event rates."),
    ("EventSounds.ts",              "Sound definitions for connect, anomaly, cluster, drilldown events."),
    ("ProceduralSoundGenerator.ts", "Web Audio API oscillator-based procedural sound generation."),
]
table(
    ["File", "Purpose"],
    audio_files,
    col_widths=[Inches(2.3), Inches(4.5)],
)

h2("12.6  Voice Control")
para("File: frontend/src/components/Voice/VoiceControl.jsx")
para(
    "Uses the browser's Web Speech API (SpeechRecognition) to capture voice input. "
    "Transcribed text is sent to POST /api/agent/voice, which routes to the T0 Agent. "
    "AgentStatusPanel.jsx shows live T0/T1 state (IDLE, LISTENING, PROCESSING, EXECUTING)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. AUTHENTICATION & SECURITY
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Authentication & Security")
divider()

auth_table = [
    ("JWT Access Token",     "Short-lived token (python-jose). Signed with JWT_SECRET_KEY env var. Verified on every protected endpoint in production mode."),
    ("Refresh Token",        "Longer-lived token. Stored in invalidated_refresh_tokens set (in-memory; Redis recommended for production)."),
    ("Password Hashing",     "bcrypt via passlib. Admin credentials stored as ADMIN_EMAIL + ADMIN_PASSWORD_HASH env vars. Plaintext passwords never stored."),
    ("Rate Limiting",        "slowapi IP-based rate limiter on /api/auth/login to prevent brute-force attacks."),
    ("SQL Injection Guard",  "validate_identifier() regex + quote_identifier() applied to all user-supplied table/column names."),
    ("CORS",                 "Strict ALLOWED_ORIGINS whitelist. Configurable per environment via .env."),
    ("No Stack Leakage",     "Global exception handler returns {error, code, path} only. Full traceback logged server-side only."),
    ("WebSocket Auth",       "In production: token query param verified via verify_token() before WebSocket accept."),
    ("Zero-Write Policy",    "Only SELECT statements ever executed against user databases. No DDL/DML."),
    ("Secret Validation",    "On startup, GOOGLE_API_KEY and JWT_SECRET_KEY are checked. Missing secrets in production mode cause RuntimeError — server won't start."),
]
table(
    ["Mechanism", "Description"],
    auth_table,
    col_widths=[Inches(1.9), Inches(4.9)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. CONFIGURATION & DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Configuration & Deployment")
divider()

h2("14.1  Environment Variables (.env)")
env_vars = [
    ("GOOGLE_API_KEY",       "Required", "Google Gemini API key for AI chat and intent classification."),
    ("JWT_SECRET_KEY",       "Required", "JWT signing secret. Must be cryptographically random in production."),
    ("ADMIN_EMAIL",          "Required", "Admin user email for login."),
    ("ADMIN_PASSWORD_HASH",  "Required", "bcrypt hash of admin password."),
    ("DB_HOST",              "Optional", "Auto-connect database host."),
    ("DB_PORT",              "Optional", "Auto-connect database port (default: 5432)."),
    ("DB_USER",              "Optional", "Auto-connect database username."),
    ("DB_PASSWORD",          "Optional", "Auto-connect database password."),
    ("DB_NAME",              "Optional", "Auto-connect database name."),
    ("PORT",                 "Optional", "Backend server port (default: 8001)."),
    ("HOST",                 "Optional", "Backend server host (default: 0.0.0.0)."),
    ("APP_ENV",              "Optional", "development (default) or production."),
    ("ALLOWED_ORIGINS",      "Optional", "Comma-separated or JSON list of CORS-allowed origins."),
    ("NEON_MAX_RETRIES",     "Optional", "Neon cold-start retry count (default: 3)."),
    ("NEON_CONNECT_TIMEOUT", "Optional", "Neon connection timeout in seconds (default: 180)."),
]
table(
    ["Variable", "Required?", "Purpose"],
    env_vars,
    col_widths=[Inches(1.9), Inches(0.8), Inches(4.1)],
)

h2("14.2  Running Locally")
numbered("python -m venv venv && venv\\Scripts\\activate  # Windows")
numbered("pip install -r backend/requirements.txt")
numbered("cp backend/.env.example backend/.env  # edit with credentials")
numbered("cd backend && python main.py  # starts on port 8001")
numbered("cd frontend && npm install && npm run dev  # starts on port 5173")

h2("14.3  Docker")
para("File: docker-compose.yml at project root.")
bullet("Defines backend and frontend services.")
bullet("Backend service: mounts backend/ volume, exposes port 8001, passes .env variables.")
bullet("Frontend service: runs npm run build → serves static from Vite preview.")
para("Start with: docker-compose up --build")

h2("14.4  Production Notes")
bullet("Set APP_ENV=production — enforces JWT auth on all routes, disables auto-reload.")
bullet("Set ALLOWED_ORIGINS to the exact frontend domain — no wildcards.")
bullet("Use a Redis-backed refresh token store instead of the in-memory set.")
bullet("Place behind a TLS-terminating reverse proxy (nginx / Caddy) — never expose uvicorn directly.")
bullet("Set ADMIN_PASSWORD_HASH to a bcrypt hash generated with: python -c \"from passlib.hash import bcrypt; print(bcrypt.hash('yourpassword'))\"")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. TEST SUITE
# ══════════════════════════════════════════════════════════════════════════════
h1("15. Test Suite")
divider()
para("Location: backend/tests/  — pytest-based test suite.")

test_files = [
    ("conftest.py",                "Pytest fixtures: mock DB connector, schema analyser, async client."),
    ("test_auth.py",               "JWT token creation, verification, expiry, and invalid token rejection."),
    ("test_db_connector.py",       "Connection pooling, identifier validation, SQL injection guard."),
    ("test_schema_analyzer.py",    "Schema introspection against mock tables and FK structures."),
    ("test_graph_generator.py",    "Node/edge generation, 3D position calculation, cluster assignment."),
    ("test_neural_core.py",        "Gravity calculation, hub score, WEZU ontology mapping."),
    ("test_anomaly_detector.py",   "Statistical outlier detection on synthetic distributions."),
    ("test_chat_service.py",       "AI chat session management, context accumulation, response format."),
    ("test_realtime_monitor.py",   "WebSocket broadcast, metrics_update payload structure."),
    ("test_advanced_logic.py",     "Integration tests for cross-service intelligence workflows."),
    ("test_safeguards.py",         "Security tests: SQL injection attempts, auth bypass, rate limit enforcement."),
    ("test_gnn.py",                "Graph Neural Network model forward pass, embedding dimension checks."),
    ("test_hybrid_fallback.py",    "Heuristic-to-NetworkX fallback when graph is too small for community detection."),
    ("ml/test_graph_core.py",      "GNN training loop, loss convergence, influence score output."),
    ("services/test_latent_manager.py", "Latent space projection consistency across connection reloads."),
]
table(
    ["File", "Test Coverage"],
    test_files,
    col_widths=[Inches(2.6), Inches(4.2)],
)
para("Run all tests: cd backend && pytest tests/ -v --tb=short")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. COMPLETE API ROUTE CATALOGUE
# ══════════════════════════════════════════════════════════════════════════════
h1("16. Complete API Route Catalogue")
divider()

all_routes = [
    # Auth
    ("POST",   "/api/auth/login",               "Authenticate user, return access + refresh tokens."),
    ("POST",   "/api/auth/refresh",             "Exchange refresh token for new access token."),
    ("POST",   "/api/auth/logout",              "Invalidate refresh token."),
    # Database
    ("POST",   "/api/database/connect",         "Connect to a database; return connection_id."),
    ("DELETE", "/api/database/disconnect/{id}", "Close connection and free pool."),
    ("GET",    "/api/database/connections",     "List all active connections."),
    # Schema
    ("POST",   "/api/schema/analyze",           "Run full schema introspection; populate NeuralCore."),
    ("GET",    "/api/schema/tables/{conn_id}",  "List all tables with metadata."),
    ("GET",    "/api/schema/columns/{conn_id}/{table}", "Column definitions for a table."),
    # Graph
    ("POST",   "/api/graph/generate",           "Generate full node-edge graph JSON with 3D coords."),
    ("GET",    "/api/graph/{conn_id}",           "Retrieve cached graph."),
    ("POST",   "/api/optimize",                 "Apply heuristic or networkx clustering."),
    # Metrics
    ("GET",    "/api/metrics/{conn_id}",        "Current TPS, row counts, vitality scores."),
    # Drilldown
    ("POST",   "/api/drilldown",                "Fetch paginated record data for a table."),
    # Hierarchy
    ("GET",    "/api/hierarchy/{conn_id}",      "Hierarchical FK flow structure."),
    # Internal Node
    ("GET",    "/api/node/{conn_id}/{table}",   "Node X-Ray: columns, indexes, sample data, health."),
    # AI
    ("POST",   "/api/ai/chat",                  "AI analyst chat message."),
    ("POST",   "/api/ai/explain",               "Natural-language explanation of a table or anomaly."),
    ("POST",   "/api/ai/classify",              "Intent classification for a user query."),
    # Agent
    ("POST",   "/api/agent/voice",              "Process voice transcription through T0 → T1 pipeline."),
    ("GET",    "/api/agent/status",             "Current T0/T1 agent state."),
    # WebSocket
    ("WS",     "/ws/{connection_id}",           "Real-time metrics stream. Token auth in production."),
    # Intelligence
    ("GET",    "/api/intelligence/anomalies",   "List detected anomalies."),
    ("GET",    "/api/intelligence/patterns",    "Mined schema/data patterns."),
    ("GET",    "/api/intelligence/predictions", "Growth and trend forecasts."),
    ("GET",    "/api/intelligence/recommendations", "Schema optimisation recommendations."),
    ("GET",    "/api/intelligence/root-cause",  "Root cause chain for a given anomaly."),
    # ML
    ("POST",   "/api/ml/analyze",              "Run scikit-learn ML analysis on a table."),
    ("POST",   "/api/ml/suggest",              "Suggest best algorithm for a table."),
    # Ontology
    ("GET",    "/api/ontology/{conn_id}",       "Retrieve WEZU domain ontology mapping."),
    # Node X-Ray
    ("GET",    "/api/node-xray/{conn_id}/{table}", "Full X-Ray: activity chart, quality radar, growth projection."),
    # Vitals
    ("GET",    "/api/vitals",                  "System vitals: CPU, memory, DB pool sizes."),
    # Simulation
    ("POST",   "/api/simulation/seed",          "Seed synthetic data into connected database."),
    # Evolution
    ("GET",    "/api/evolution/snapshots/{id}", "List schema evolution snapshots."),
    ("POST",   "/api/evolution/replay",         "Start evolution timeline playback."),
    # Health
    ("GET",    "/health",                       "Basic health check: {status: ok, version}."),
    ("GET",    "/health/ready",                 "Readiness check including DB connectivity."),
    ("GET",    "/health/routers",               "Router registry status report."),
]
table(
    ["Method", "Path", "Description"],
    all_routes,
    col_widths=[Inches(0.7), Inches(2.7), Inches(3.4)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. FUTURE ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
h1("17. Future Roadmap")
divider()

roadmap = [
    ("v2.0", "Predictive 3D Simulation",
     "ML forecasts projected back onto Valkyrie 3D graph. Nodes change size and position to show 'Predicted Future States' based on multi-day ML model output."),
    ("v2.1", "Cross-Table Feature Joins",
     "Work on Data will perform automatic SQL JOINs on FK chains, enriching the ML feature matrix with relational context from multiple tables."),
    ("v2.2", "Model Export & Scheduling",
     "Trained scikit-learn models (pickle) and analysis summaries (JSON) become exportable. Scheduled re-runs (cron) enable automated daily model refresh."),
    ("v2.3", "AutoML with Optuna",
     "Bayesian hyperparameter optimisation integrated into the ML pipeline. Automatically tunes Random Forest, Gradient Boosting, and clustering parameters within a configurable time budget."),
    ("v2.4", "Redis-Backed State",
     "Migrate in-memory token store, agent state, and schema snapshots to Redis for horizontal scalability and multi-instance deployments."),
    ("v2.5", "Graph Neural Network (GNN) Production",
     "Promote the GNN model (backend/ml/gnn_model.py) from research to production for link prediction and anomaly scoring on the schema graph."),
    ("v3.0", "LLM-Native Query Interface",
     "Plain English analytical requests ('find customers most likely to churn') automatically translated to algorithm family, column selection, and model configuration."),
    ("v3.1", "Multi-Tenant SaaS",
     "Per-tenant schema isolation, connection quotas, usage billing, and organisation-level audit logging."),
    ("v3.2", "Streaming Ingestion",
     "Kafka / Pulsar connector to ingest real-time event streams, visualised as live particle flows on the 3D graph without database polling."),
]
table(
    ["Version", "Feature", "Description"],
    roadmap,
    col_widths=[Inches(0.6), Inches(1.9), Inches(4.3)],
)

doc.add_paragraph()
divider()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    "Living Data Intelligence Platform  ·  Comprehensive Architectural & Functional Documentation  ·  Version 1.0  ·  March 2026"
)
set_font(r, size=9, italic=True, color=C_MID_GRAY)

# ─── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\karth\living-data-intelligence-backend-sasir\COMPREHENSIVE_ARCHITECTURE_DOCUMENTATION.docx"
doc.save(out)
print(f"Saved: {out}")
