"""
Generate Work on Data — Complete Word Documentation
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document()

# ─── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ─── Style helpers ────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p

def add_para(text="", bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.35)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        set_font(run, bold=True, color=(0, 70, 127))
    run = p.add_run(text)
    set_font(run)
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        set_font(run, bold=True)
    run = p.add_run(text)
    set_font(run)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x3a)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "004680")
        tcPr.append(shd)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for run in row_cells[i].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = col_widths[j]
    return table

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "004680")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("WORK ON DATA")
set_font(run, size=28, bold=True, color=(0, 70, 127))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Complete Technical & Functional Documentation")
set_font(run, size=16, color=(80, 80, 80))

doc.add_paragraph()
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run("Living Data Intelligence Platform  ·  Version 1.1")
set_font(run, size=12, italic=True, color=(120, 120, 120))

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("March 2026")
set_font(run, size=11, italic=True, color=(150, 150, 150))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("Table of Contents", level=1)
toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "System Architecture"),
    ("3.", "The Four Analytical Pillars"),
    ("  3.1", "Classification Family"),
    ("  3.2", "Regression Family"),
    ("  3.3", "Time Series Family"),
    ("  3.4", "Clustering Family"),
    ("4.", "The Intelligence Pipeline: From SQL to Tensors"),
    ("  4.1", "Stage 1 — Secure Data Ingestion"),
    ("  4.2", "Stage 2 — Feature Engineering & Preprocessing"),
    ("  4.3", "Stage 3 — Multi-Threaded Execution"),
    ("  4.4", "Stage 4 — Insight Generation"),
    ("5.", "Frontend: Work on Data Modal"),
    ("  5.1", "AI Suggestion Engine"),
    ("  5.2", "Algorithm Selection Interface"),
    ("  5.3", "Results Panel"),
    ("6.", "Frontend: Deep Analysis Page"),
    ("  6.1", "2D Scatter & Relationship Plots"),
    ("  6.2", "3D Point Cloud (WebGL)"),
    ("  6.3", "AI Analyst Chat"),
    ("7.", "Backend: API Reference"),
    ("  7.1", "POST /api/ml/analyze"),
    ("  7.2", "POST /api/ml/suggest"),
    ("8.", "Mathematical Foundations"),
    ("9.", "AI Insights Engine"),
    ("10.", "Operational Safety & Read-Only Policy"),
    ("11.", "Component Mapping"),
    ("12.", "Future Roadmap"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    set_font(run, size=11, color=(0, 70, 127) if not num.startswith("  ") else (60, 60, 60))
    if not num.startswith("  "):
        run.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. Executive Summary", level=1)
divider()

add_para(
    "The Work on Data subsystem is the Living Data Intelligence Platform's primary bridge "
    "between raw database storage and predictive intelligence. While the main 3D graph "
    "(Valkyrie Engine) serves to visualise the topological relationships of a database, "
    "the Work on Data module is dedicated to the statistical and predictive relationships "
    "within the rows themselves."
)
doc.add_paragraph()

add_para("It is a non-invasive, read-only analytical workbench that allows users to:")
add_bullet("Extract data samples securely from any connected database.", bold_prefix="Extract:")
add_bullet("Transform heterogeneous database types into standardised ML-ready tensors.", bold_prefix="Transform:")
add_bullet("Execute state-of-the-art algorithms across four analytical families.", bold_prefix="Execute:")
add_bullet("Visualise results in a high-fidelity Deep Analysis environment.", bold_prefix="Visualise:")
add_bullet("Interpret findings via a context-aware AI Analyst Chat.", bold_prefix="Interpret:")
doc.add_paragraph()

add_para(
    "The platform operates under a strict Zero-Write Policy: every SQL statement is a "
    "SELECT, all computations happen in server RAM, and no data is ever written back to "
    "the user's database. This guarantees zero footprint on production OLTP performance."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. System Architecture", level=1)
divider()

add_para(
    "The platform follows a three-layer architecture: a React analytical frontend, a "
    "FastAPI ML backend engine, and an AI/LLM interpretation layer."
)
doc.add_paragraph()

add_heading("Architecture Layers", level=2)
layers = [
    ("Frontend: React Analytical Layer",
     "WorkOnDataModal, DeepAnalysisPage, AI Suggestion Engine, Recharts + React-Three-Fiber visualisers, AI Analyst Chat"),
    ("Backend: FastAPI ML Engine",
     "POST /api/ml/analyze, db_connector (SQL SELECT), Preprocessing Engine, scikit-learn Executor (thread pool), AnalysisResult JSON serialiser"),
    ("AI / LLM Layer",
     "Context-aware prompt builder, Low-temperature (0.2) inference, Natural language explanations of statistical output"),
]
for title, detail in layers:
    add_para(detail, bold_prefix=title + ":")
    doc.add_paragraph()

add_heading("Data Flow", level=2)
steps = [
    "User opens WorkOnDataModal → frontend AI heuristic runs schema scoring.",
    "Optional: POST /api/ml/suggest enhances the recommendation with server-side schema metadata.",
    "User accepts suggestion or configures manually, then clicks Run Model or Open Deep Analysis.",
    "Frontend POSTs to /api/ml/analyze with {connection_id, table, family, algo, target, features}.",
    "Backend fetches ≤5,000 rows via safe-quoted SELECT, preprocesses, trains in a background thread.",
    "AnalysisResult JSON (metrics, feature importances, predictions, insights, scatter sample) returned.",
    "Frontend renders ResultsPanel inline, or DeepAnalysisPage in a new tab.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(step)
    set_font(run)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE FOUR ANALYTICAL PILLARS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. The Four Analytical Pillars", level=1)
divider()

add_para(
    "The platform categorises all data-science workflows into four mutually exclusive "
    "Algorithm Families. Each family targets a distinct analytical objective and ships "
    "with dedicated preprocessing, evaluation metrics, and insight logic."
)

# ── 3.1 Classification ────────────────────────────────────────────────────────
add_heading("3.1  Classification Family  (Supervised)", level=2)

add_para(
    "Predict a discrete category (label) based on a set of independent features. "
    "Classification is appropriate when the target column contains distinct classes such as "
    "product category, customer segment, or fraud indicator."
)
doc.add_paragraph()

add_heading("Supported Algorithms", level=3)
clf_table = [
    ("rf_clf",  "Random Forest",       "100 decision trees; Gini impurity criterion; max_depth=10; n_jobs=-1",
     "Mixed types (Numeric + Categorical); high-dimensional data"),
    ("svm",     "Support Vector Machine", "RBF kernel; C=1.0; probability=True; StandardScaler applied",
     "Clear margin separation; small-medium datasets"),
    ("knn",     "K-Nearest Neighbors", "k = min(5, n-1); StandardScaler applied",
     "Local patterns; small-medium datasets"),
    ("logreg",  "Logistic Regression", "lbfgs solver; max_iter=1000; C=1.0; multi_class=auto; StandardScaler applied",
     "Binary/multi-class; interpretable coefficients; linear boundaries"),
]
add_table(
    ["ID", "Name", "Implementation Detail", "Best For"],
    clf_table,
    col_widths=[Inches(0.9), Inches(1.4), Inches(2.5), Inches(2.0)],
)
doc.add_paragraph()

add_heading("Evaluation Metrics", level=3)
clf_metrics = [
    ("Accuracy",   "Ratio of correct predictions to total predictions."),
    ("F1-Score",   "Harmonic mean of Precision and Recall. Used as primary quality indicator."),
    ("Precision",  "Accuracy of positive predictions (TP / (TP + FP))."),
    ("Recall",     "Ability to find all positive samples (TP / (TP + FN))."),
    ("n_classes",  "Number of distinct label classes detected in the target column."),
    ("train_size", "Number of rows used for training (80% stratified split)."),
    ("test_size",  "Number of rows used for evaluation (20% stratified split)."),
]
for metric, desc in clf_metrics:
    add_bullet(desc, bold_prefix=metric + ":")

doc.add_paragraph()
add_heading("AI Insight Logic — Classification", level=3)
add_para("The insight engine emits the following natural-language observations:")
add_bullet("Model accuracy and F1-score on held-out test set with class count.")
add_bullet("Most discriminative feature and its percentage importance.")
add_bullet("Generalisation quality judgement: Excellent (F1 ≥ 0.85), Good (F1 ≥ 0.70), or Moderate.")
add_bullet("Precision vs. Recall imbalance warning with actionable threshold advice.")

divider()

# ── 3.2 Regression ────────────────────────────────────────────────────────────
add_heading("3.2  Regression Family  (Supervised)", level=2)

add_para(
    "Predict a continuous numeric value (quantity) based on independent features. "
    "Regression is appropriate when the target column is a numeric measurement such as "
    "revenue, price, quantity, or response time."
)
doc.add_paragraph()

add_heading("Supported Algorithms", level=3)
reg_table = [
    ("linear",  "Linear Regression",    "Ordinary Least Squares; no scaling required; coef_ used as importance",
     "Linear relationships; baseline model"),
    ("ridge",   "Ridge Regression",     "L2 regularisation; alpha=1.0; StandardScaler applied",
     "Multicollinearity; many features"),
    ("lasso",   "Lasso Regression",     "L1 regularisation; alpha=0.1; max_iter=5000; StandardScaler applied",
     "Feature selection; sparse solutions"),
    ("xgboost", "Gradient Boosting",    "GradientBoostingRegressor; n_estimators=100; max_depth=4; lr=0.1; subsample=0.8",
     "Tabular data; non-linear; highest accuracy"),
]
add_table(
    ["ID", "Name", "Implementation Detail", "Best For"],
    reg_table,
    col_widths=[Inches(0.9), Inches(1.4), Inches(2.5), Inches(2.0)],
)
doc.add_paragraph()

add_heading("Evaluation Metrics", level=3)
reg_metrics = [
    ("R²",   "Proportion of variance in the target explained by the model (0.0 – 1.0; clipped at -1.0 lower bound)."),
    ("RMSE", "Root Mean Square Error — standard deviation of residuals; penalises large errors heavily."),
    ("MAE",  "Mean Absolute Error — average absolute difference between predicted and actual; robust to outliers."),
]
for metric, desc in reg_metrics:
    add_bullet(desc, bold_prefix=metric + ":")

doc.add_paragraph()
add_heading("Prediction Output", level=3)
add_para(
    "After training, the engine produces a 6-period forward extrapolation "
    "(Period +1 through Quarter) based on the test-set mean prediction with "
    "linearly expanding confidence intervals. Each prediction carries a "
    "confidence label: high (periods 1–2), medium (3–4), or low (5–6)."
)

divider()

# ── 3.3 Time Series ────────────────────────────────────────────────────────────
add_heading("3.3  Time Series Family  (Temporal Forecasting)", level=2)

add_para(
    "Predict future values based on past chronological observations. The platform uses a "
    "custom Harmonic Regression model requiring no external dependencies (no statsmodels, "
    "no Prophet library required at runtime)."
)
doc.add_paragraph()

add_heading("Algorithm: Trend + Seasonal Harmonic Regression", level=3)
add_para(
    "The time-series engine decomposes each series into two additive components:"
)
add_bullet(
    "Linear Trend — fitted via NumPy polyfit (degree 1). "
    "Captures long-term directional movement (slope and intercept).",
    bold_prefix="Component 1:"
)
add_bullet(
    "Weekly Seasonal Periodicity — fitted via harmonic regression on trend residuals. "
    "Up to k=3 harmonics with period T=7 (weekly). "
    "Uses np.linalg.lstsq to solve sinusoidal coefficients.",
    bold_prefix="Component 2:"
)
doc.add_paragraph()

add_heading("30-Day Forecast", level=3)
add_para(
    "The fitted trend and seasonal components are extrapolated 30 days into the future. "
    "Confidence intervals are proportional to the magnitude of each forecast point "
    "and widen with forecast horizon (5% base, expanding by 15% per step)."
)
doc.add_paragraph()

add_heading("Evaluation Metrics", level=3)
ts_metrics = [
    ("MAPE",           "Mean Absolute Percentage Error on last 20% holdout — primary accuracy indicator."),
    ("RMSE",           "Root Mean Square Error on holdout set."),
    ("MAE",            "Mean Absolute Error on holdout set."),
    ("trend",          "Direction: upward, downward, or flat (based on linear slope sign)."),
    ("monthly_growth", "Percentage change per 30 days implied by the fitted trend slope."),
]
for metric, desc in ts_metrics:
    add_bullet(desc, bold_prefix=metric + ":")

add_para("Date-column auto-detection: the engine scans feature_cols first, then all dataframe columns for datetime dtype.")

divider()

# ── 3.4 Clustering ────────────────────────────────────────────────────────────
add_heading("3.4  Clustering Family  (Unsupervised)", level=2)

add_para(
    "Group similar records together based on spatial proximity, without a predefined target. "
    "Clustering reveals natural segments — e.g. customer behavioural groups, anomaly clusters, "
    "or product similarity buckets — with no label column required."
)
doc.add_paragraph()

add_heading("Supported Algorithms", level=3)
cl_table = [
    ("kmeans", "K-Means",
     "Auto-K selection (k=2–8) via Silhouette Score; n_init=10; StandardScaler applied",
     "Well-separated spherical clusters"),
    ("dbscan", "DBSCAN",
     "Auto-eps via k-distance 90th percentile (k=5); min_samples=5; StandardScaler applied",
     "Arbitrary cluster shapes; noise/outlier detection"),
]
add_table(
    ["ID", "Name", "Implementation Detail", "Best For"],
    cl_table,
    col_widths=[Inches(0.9), Inches(1.2), Inches(2.7), Inches(2.0)],
)
doc.add_paragraph()

add_heading("K-Means Auto-K Selection", level=3)
add_para(
    "The platform automatically selects the optimal number of clusters by iterating "
    "k from 2 to min(8, n // 10) and retaining the model with the highest Silhouette Score. "
    "sample_size=min(2000, n) is used for the Silhouette calculation to keep runtime bounded."
)
doc.add_paragraph()

add_heading("DBSCAN Auto-eps Tuning", level=3)
add_para(
    "Epsilon is set to the 90th percentile of the k-NN distance distribution (k=5), "
    "which is a standard heuristic for detecting the 'elbow' in the k-distance plot."
)
doc.add_paragraph()

add_heading("Evaluation Metrics", level=3)
cl_metrics = [
    ("n_clusters",       "Number of clusters discovered (DBSCAN excludes noise label -1)."),
    ("silhouette_score", "Mean silhouette coefficient — measure of intra-cluster cohesion vs. inter-cluster separation (max 1.0)."),
    ("inertia",          "K-Means only — sum of squared distances of samples to their nearest cluster centre."),
    ("n_noise_points",   "DBSCAN only — count of samples labelled as noise (label = -1)."),
]
for metric, desc in cl_metrics:
    add_bullet(desc, bold_prefix=metric + ":")

add_heading("Feature Importance in Clustering", level=3)
add_para(
    "Because clustering is unsupervised (no target), feature importance is derived from "
    "centroid separation: the standard deviation of cluster centroids along each feature axis, "
    "normalised to sum to 1.0. Features with higher centroid spread are more discriminating."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — THE INTELLIGENCE PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. The Intelligence Pipeline: From SQL to Tensors", level=1)
divider()

add_para(
    "The transformation process inside backend/app/api/ml_analysis.py is a four-stage "
    "assembly line that converts raw database rows into trained models and insights."
)
doc.add_paragraph()

# Stage 1
add_heading("4.1  Stage 1 — Secure Data Ingestion", level=2)
add_para(
    "A SELECT query is constructed using the _safe_quote helper, which applies "
    "database-appropriate identifier quoting:"
)
add_bullet("MySQL: backtick quoting — `table_name`, `column_name`")
add_bullet("PostgreSQL / others: double-quote quoting — \"table_name\", \"column_name\"")
add_para("Additional constraints applied at ingestion:")
add_bullet("Row cap: LIMIT is enforced in the range [100, 5000]. Default sampling is 2,000 rows.")
add_bullet("Column cap: Maximum 20 columns are fetched in a single query.")
add_bullet("Only SELECT statements are generated — no INSERT, UPDATE, DELETE, or DDL.")
doc.add_paragraph()

add_code_block(
    "query = f\"SELECT {safe_cols} FROM {safe_table} LIMIT {limit}\"\n"
    "# safe_cols and safe_table are produced by _safe_quote(), never from user strings directly."
)
doc.add_paragraph()

# Stage 2
add_heading("4.2  Stage 2 — Feature Engineering & Preprocessing", level=2)
add_para("Raw database data is often 'dirty.' The _preprocess function applies automatic corrections:")

preprocessing_steps = [
    ("Row filtering",     "Rows where ALL selected feature columns are null are dropped."),
    ("Numeric imputation","Missing numeric values are filled with the column median. "
                          "If median is NaN, zero is used as fallback."),
    ("Categorical impute","Missing string/object values are filled with the sentinel '__missing__'."),
    ("Label encoding",    "Categorical strings are integer-encoded via scikit-learn LabelEncoder."),
    ("NaN sanitisation",  "After hstacking, any residual NaN/±Inf values are replaced with 0.0 via np.nan_to_num."),
    ("Target encoding",   "Classification targets: LabelEncoder. "
                          "Regression/timeseries targets: numeric, median-imputed; non-numeric coerced."),
    ("Tensor packing",    "All features are hstacked into a 2D NumPy ndarray X. Target is a 1D ndarray y."),
]
add_table(
    ["Step", "Description"],
    preprocessing_steps,
    col_widths=[Inches(1.6), Inches(5.2)],
)
doc.add_paragraph()

add_para("Minimum data requirement: 10 rows minimum for clustering; 20 rows for supervised learning.")
doc.add_paragraph()

# Stage 3
add_heading("4.3  Stage 3 — Multi-Threaded Execution", level=2)
add_para(
    "To prevent a CPU-intensive scikit-learn training task from blocking the FastAPI async event loop, "
    "the platform uses asyncio.get_running_loop().run_in_executor(None, fn, *args). "
    "This offloads computation to Python's default ThreadPoolExecutor (one thread per logical CPU), "
    "keeping the web server responsive to other requests during training."
)
doc.add_paragraph()

add_code_block(
    "loop = asyncio.get_running_loop()\n"
    "metrics, fi, predictions = await loop.run_in_executor(\n"
    "    None, _run_classification, X, y, req.algo, feature_names\n"
    ")"
)
doc.add_paragraph()

# Stage 4
add_heading("4.4  Stage 4 — Insight Generation", level=2)
add_para(
    "Once training completes, the _build_insights function converts raw numeric metrics "
    "into a list of human-readable strings. The insight logic is family-specific:"
)
insight_families = [
    ("Classification", "Accuracy + F1 summary; top feature; generalisation quality; precision/recall imbalance warning."),
    ("Regression",     "R² variance explanation; RMSE + MAE summary; top two predictors; fit quality judgement."),
    ("Time Series",    "Trend direction + monthly growth rate; MAPE quality band (<10%, 10–25%, >25%)."),
    ("Clustering",     "Cluster count + silhouette score; top separating feature; cluster strength; DBSCAN noise audit."),
]
for family, logic in insight_families:
    add_bullet(logic, bold_prefix=family + ":")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — FRONTEND: WORK ON DATA MODAL
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. Frontend: Work on Data Modal", level=1)
divider()

add_para(
    "The WorkOnDataModal component (frontend/src/components/Dashboard/WorkOnDataModal.jsx) "
    "is a full-screen overlay launched from the main dashboard. It provides a guided "
    "configuration experience backed by an AI recommendation engine."
)
doc.add_paragraph()

# 5.1
add_heading("5.1  AI Suggestion Engine", level=2)
add_para(
    "On modal open, a two-phase suggestion process runs automatically:"
)

add_heading("Phase 1 — Client-side heuristic (instant)", level=3)
add_para(
    "The generateAISuggestion function scores all graph nodes using:"
)
add_code_block(
    "score = log10(row_count + 1) * 40\n"
    "       + column_count * 2\n"
    "       + foreign_key_count * 5"
)
add_para("The highest-scoring table is selected as primary. Algorithm family is chosen by:")
add_bullet("Time Series (82% confidence) — if date/timestamp columns AND numeric columns exist.")
add_bullet("Regression, XGBoost (91%) — if ≥2 numeric columns exist.")
add_bullet("Classification, Random Forest (85%) — if categorical target columns exist.")
add_bullet("Clustering, K-Means (74%) — fallback when no clear target is detectable.")
doc.add_paragraph()

add_heading("Phase 2 — Server-side schema enrichment (async)", level=3)
add_para(
    "A POST request to /api/ml/suggest is fired with the primary table name. "
    "The server uses schema_analyzer metadata (column types, row counts) to validate "
    "and potentially override the client suggestion. The enriched result is tagged "
    "'Schema-enhanced' in the UI."
)
doc.add_paragraph()

# 5.2
add_heading("5.2  Algorithm Selection Interface", level=2)
add_para(
    "The right panel presents a collapsible accordion of the four algorithm families. "
    "Each family shows its algorithms with:"
)
add_bullet("Tag badge — quality descriptor: Accurate, Robust, Fast, Interpretable, etc.")
add_bullet("Best For — one-line description of the ideal use case.")
add_bullet("Active highlighting — selected family/algorithm is highlighted with family accent colour.")

add_para("Column configuration (left panel):")
add_bullet("Primary Table — sorted by AI score; maximum 12 shown; AI-selected table is tagged 'AI'.")
add_bullet("Join Tables — optional secondary tables (checkboxes); AI-suggested joins show reasoning.")
add_bullet("Target Column — toggle buttons with type labels (NUM / DATE / CAT) for fast identification.")
add_bullet("Feature Columns — toggle buttons; AI-recommended columns marked with a star (★).")
doc.add_paragraph()

# 5.3
add_heading("5.3  Results Panel", level=2)
add_para("After clicking Run Model, the ResultsPanel renders inline with the following cards:")

results_cards = [
    ("Metrics Row",        "Grid of key/value metric cards rendered from the backend response. Excludes internal keys (samples, train_size, test_size, n_classes, model)."),
    ("Feature Importance", "Animated horizontal bar chart (framer-motion) showing up to 8 features sorted descending. Bar colour tied to family accent."),
    ("Predictions Panel",  "For clustering/classification: segment/class distribution as percentage bars. For regression/timeseries: period labels with confidence-coded colours."),
    ("AI Insights",        "Numbered list of natural-language observations generated by _build_insights, animated with staggered fade-in."),
]
add_table(
    ["Card", "Content"],
    results_cards,
    col_widths=[Inches(1.8), Inches(5.0)],
)
doc.add_paragraph()

add_para(
    "An Open Deep Analysis button in the results header constructs a URL with all "
    "configuration parameters serialised as query-string parameters and opens "
    "/deep-analysis in a new browser tab."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — FRONTEND: DEEP ANALYSIS PAGE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. Frontend: Deep Analysis Page  (/deep-analysis)", level=1)
divider()

add_para(
    "The DeepAnalysisPage component (frontend/src/components/Dashboard/DeepAnalysisPage.jsx) "
    "is a full-screen, standalone research dashboard optimised for intensive data exploration. "
    "It reads its configuration from URL query-string parameters."
)
doc.add_paragraph()

add_heading("6.1  2D Scatter & Relationship Plots", level=2)
add_para("Rendered using Recharts:")
add_bullet("Cluster/class coloured scatter points on X–Y axes.")
add_bullet("Interactive tooltips reveal raw database row values on hover.")
add_bullet("Feature Importance: dynamic bar chart — which columns drive the model's decisions.")
add_bullet("Trend & Forecast: composed chart showing actual vs. predicted values with 95% confidence interval bands.")
add_bullet("Correlation Heatmap: matrix visualisation of inter-feature Pearson correlations (Cyan = Positive, Purple = Negative).")
doc.add_paragraph()

add_heading("6.2  3D Point Cloud  (WebGL)", level=2)
add_para("Rendered using React-Three-Fiber (Three.js):")
add_bullet("Up to 80 scatter_sample rows projected onto X, Y, Z axes.")
add_bullet("Points coloured by cluster ID or classification label.")
add_bullet("Full orbit/zoom/pan navigation — users can 'fly through' their data.")
add_bullet("Useful for verifying whether clusters are well-separated or overlapping in 3D feature space.")
doc.add_paragraph()

add_heading("6.3  AI Analyst Chat  (The 'Brain')", level=2)
add_para(
    "A persistent AI chat panel is embedded in the Deep Analysis page. "
    "Unlike a generic assistant, it is primed with the full analysis context:"
)
add_bullet("Temperature: 0.2 — deterministic, analytical responses.")
add_bullet("System context: raw scikit-learn metrics + feature importance array + prediction series.")
add_bullet("User can ask questions such as: 'What are the primary drivers for this target?' or 'Is my model biased?'")
add_bullet("The AI generates Natural Language Interface to Statistics — no statistical knowledge required.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("7. Backend: API Reference", level=1)
divider()

add_heading("7.1  POST /api/ml/analyze", level=2)
add_para("The core ML execution endpoint. All computation is performed here.")
doc.add_paragraph()

add_heading("Request Body (AnalysisRequest)", level=3)
req_fields = [
    ("connection_id",    "str",                "Required", "Active database connection identifier."),
    ("table",            "str",                "Required", "Primary table name to analyse."),
    ("secondary_tables", "List[str]",          "Optional", "Additional tables to join (currently informational)."),
    ("family",           "str",                "Required", "One of: classification | regression | timeseries | clustering"),
    ("algo",             "str",                "Required", "Algorithm ID: rf_clf | svm | knn | logreg | linear | ridge | lasso | xgboost | arima | prophet | kmeans | dbscan"),
    ("target",           "Optional[str]",      "Optional", "Target column name (required for supervised families)."),
    ("features",         "Optional[List[str]]","Optional", "Feature column names. Falls back to all non-target columns (max 10) if empty."),
]
add_table(
    ["Field", "Type", "Required", "Description"],
    req_fields,
    col_widths=[Inches(1.4), Inches(1.2), Inches(0.8), Inches(3.4)],
)
doc.add_paragraph()

add_heading("Response Body (AnalysisResult)", level=3)
resp_fields = [
    ("algo",                "str",             "Algorithm ID echoed from request."),
    ("family",              "str",             "Algorithm family echoed from request."),
    ("table",               "str",             "Table name echoed from request."),
    ("row_count",           "int",             "Rows fetched (or schema count if fetch returned 0)."),
    ("metrics",             "Dict[str, Any]",  "Family-specific numeric metrics (see Section 3)."),
    ("feature_importances", "List[{name, importance}]", "Sorted descending by importance (0.0–1.0)."),
    ("predictions",         "List[Prediction]","Up to 6 items: {label, value, lower, upper, confidence}."),
    ("insights",            "List[str]",       "Natural-language insight strings."),
    ("scatter_sample",      "List[Dict]",      "Up to 80 raw rows (JSON-serialisable values only) for 2D/3D plots."),
    ("status",              "str",             "Always 'success' on 200 response."),
]
add_table(
    ["Field", "Type", "Description"],
    resp_fields,
    col_widths=[Inches(1.8), Inches(1.6), Inches(3.4)],
)
doc.add_paragraph()

add_heading("Error Responses", level=3)
add_bullet("422 Unprocessable Entity — validation errors (insufficient data, missing target, < 2 classes, etc.).")
add_bullet("500 Internal Server Error — unexpected scikit-learn or database failure (full traceback logged server-side).")
doc.add_paragraph()

add_heading("7.2  POST /api/ml/suggest", level=2)
add_para(
    "Heuristic recommender. Returns the best algorithm family, algorithm, target, and "
    "feature columns for a given table based on its schema metadata."
)
doc.add_paragraph()

add_heading("Query Parameters", level=3)
add_bullet("connection_id (str) — Active connection identifier.")
add_bullet("table (str) — Table name to inspect.")
doc.add_paragraph()

add_heading("Response Structure", level=3)
add_code_block(
    "{\n"
    "  \"suggestion\": {\n"
    "    \"family\":     \"regression\",\n"
    "    \"algo\":       \"xgboost\",\n"
    "    \"target\":     \"revenue\",\n"
    "    \"features\":   [\"quantity\", \"price\", \"category\"],\n"
    "    \"confidence\": 91\n"
    "  },\n"
    "  \"reason\": \"'orders' has 12 numeric columns and 45,230 rows. GradientBoosting excels on tabular data.\",\n"
    "  \"table_info\": { \"row_count\": 45230, \"column_count\": 15 }\n"
    "}"
)
doc.add_paragraph()

add_heading("Suggestion Priority Rules", level=3)
suggest_rules = [
    ("1st priority", "timeseries (83%)",     "Table has date/timestamp column AND numeric columns."),
    ("2nd priority", "regression (91%)",     "Table has ≥ 2 numeric columns."),
    ("3rd priority", "classification (85%)", "Table has categorical (varchar/text/bool/enum) columns."),
    ("Fallback",     "clustering (74%)",     "None of the above conditions are met."),
]
add_table(
    ["Priority", "Family (Confidence)", "Condition"],
    suggest_rules,
    col_widths=[Inches(1.0), Inches(1.7), Inches(4.1)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — MATHEMATICAL FOUNDATIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8. Mathematical Foundations", level=1)
divider()

add_heading("8.1  Feature Importance for Tree-Based Models", level=2)
add_para(
    "For Random Forest and Gradient Boosting, feature importance is the total reduction "
    "of the impurity criterion (Gini or MSE) brought by that feature across all nodes "
    "in all trees, normalised to sum to 1.0:"
)
add_code_block(
    "Importance(f) = SUM over all nodes n that split on f of:\n"
    "    (N_n / N) * (Gain_n - N_nL/N_n * Gain_nL - N_nR/N_n * Gain_nR)\n\n"
    "Where: N = total samples, N_n = samples at node n,\n"
    "       N_nL / N_nR = samples in left/right child,\n"
    "       Gain = Gini impurity (classification) or MSE (regression)"
)
doc.add_paragraph()

add_heading("8.2  Logistic Regression Importance", level=2)
add_para(
    "For logistic regression, feature importance is derived from the absolute values of "
    "the model coefficients, averaged across classes for multi-class problems:"
)
add_code_block(
    "importance_i = mean(|coef_class_j_i|) for all classes j"
)
doc.add_paragraph()

add_heading("8.3  Seasonal Harmonic Decomposition", level=2)
add_para(
    "For time series, the platform fits a linear trend first, then models the residuals "
    "as a sum of sinusoids (Fourier series) to capture seasonal periodicity:"
)
add_code_block(
    "y = (slope * t + intercept)          # Linear trend\n"
    "  + SUM_{k=1}^{K} [\n"
    "        a_k * cos(2*pi*k*t / T)\n"
    "      + b_k * sin(2*pi*k*t / T)\n"
    "    ]                                # Seasonal component\n\n"
    "Where: T = 7 (weekly period), K = min(3, n//14)\n"
    "Coefficients a_k, b_k solved via np.linalg.lstsq on residuals."
)
doc.add_paragraph()

add_heading("8.4  Silhouette Score", level=2)
add_para("Measures cluster quality for any clustering algorithm:")
add_code_block(
    "s(i) = (b(i) - a(i)) / max(a(i), b(i))\n\n"
    "Where: a(i) = mean intra-cluster distance for point i\n"
    "       b(i) = mean distance to nearest neighbouring cluster\n"
    "Range: [-1, 1]  (higher is better; >0.5 = strong; 0.25–0.5 = moderate)"
)
doc.add_paragraph()

add_heading("8.5  MAPE (Time Series Accuracy)", level=2)
add_code_block(
    "MAPE = (1/n) * SUM |( y_i - y_hat_i ) / y_i| * 100\n"
    "(Computed only over non-zero actuals; capped at 9999.0 to prevent division anomalies)"
)
doc.add_paragraph()

add_heading("8.6  DBSCAN Auto-eps", level=2)
add_code_block(
    "eps = 90th_percentile( kNN_distances(X_scaled, k=5) )\n"
    "eps = max(eps, 0.1)   # floor to prevent degenerate single-cluster"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — AI INSIGHTS ENGINE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("9. AI Insights Engine", level=1)
divider()

add_para(
    "The _build_insights function translates quantitative model output into actionable "
    "natural-language observations. It runs server-side, immediately after training, "
    "and returns a list of strings included in the AnalysisResult payload."
)
doc.add_paragraph()

insight_table = [
    ("Classification",
     "F1 ≥ 0.85",
     "Excellent generalisation — model is ready for production scoring."),
    ("Classification",
     "F1 < 0.70",
     "Moderate performance. More labelled data or feature engineering may improve F1."),
    ("Classification",
     "Precision > Recall by >8%",
     "Model is conservative. Lower decision threshold for higher recall if false-negatives are costly."),
    ("Regression",
     "R² ≥ 0.85",
     "Strong predictive power — suitable for production forecasting."),
    ("Regression",
     "R² < 0.55",
     "Low R² suggests high noise or non-linear patterns. Consider GradientBoosting / feature engineering."),
    ("Time Series",
     "MAPE < 10%",
     "Low MAPE (<10%) — forecasts are highly reliable for short-term planning."),
    ("Time Series",
     "MAPE > 25%",
     "High MAPE (>25%) — series is volatile. Use prediction intervals; do not rely on point forecasts alone."),
    ("Clustering",
     "Silhouette ≥ 0.5",
     "Strong cluster structure — segments are well-separated and actionable."),
    ("Clustering",
     "Silhouette < 0.25",
     "Weak cluster structure. Data may be continuous; consider reducing k."),
    ("Clustering (DBSCAN)",
     "noise_points > 0",
     "DBSCAN flagged {n} noise points ({pct}%) as outliers — review for data quality issues."),
]
add_table(
    ["Family", "Condition", "Generated Insight"],
    insight_table,
    col_widths=[Inches(1.2), Inches(1.7), Inches(3.9)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — OPERATIONAL SAFETY
# ══════════════════════════════════════════════════════════════════════════════

add_heading("10. Operational Safety & Read-Only Policy", level=1)
divider()

add_para("The platform strictly adheres to a Zero-Write Policy at every layer:")
doc.add_paragraph()

safety_rules = [
    ("SQL Isolation",      "Only SELECT statements are executed. The _safe_quote helper prevents identifier injection. LIMIT is always applied."),
    ("No DDL / DML",       "No INSERT, UPDATE, DELETE, CREATE, DROP, or TRUNCATE statements are ever generated."),
    ("No Temporary Tables","All analysis is performed entirely in server RAM (pandas DataFrames and NumPy arrays). Nothing is spilled back to the database."),
    ("Low Impact",         "Row-capping at ≤ 5,000 rows ensures that even complex clustering tasks complete in under 2 seconds on typical hardware."),
    ("Memory Safety",      "np.nan_to_num and median/mode imputation prevent crashes on dirty data. Minimum row counts (≥10 or ≥20) prevent degenerate model training."),
    ("Thread Safety",      "ML training runs in a ThreadPoolExecutor via run_in_executor, preventing it from blocking the FastAPI async event loop."),
    ("Error Isolation",    "ValueError (user-facing data issues) → HTTP 422. All other exceptions → HTTP 500 with server-side traceback logging. No internal details exposed to clients."),
]
add_table(
    ["Safeguard", "Description"],
    safety_rules,
    col_widths=[Inches(1.6), Inches(5.2)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — COMPONENT MAPPING
# ══════════════════════════════════════════════════════════════════════════════

add_heading("11. Component Mapping", level=1)
divider()

component_map = [
    ("backend/app/api/ml_analysis.py",
     "FastAPI router; Pydantic models; data fetching; preprocessing; "
     "classification/regression/clustering/timeseries runners; insight builder; "
     "/analyze and /suggest endpoints."),
    ("frontend/src/components/Dashboard/WorkOnDataModal.jsx",
     "ML launcher modal; client-side AI suggestion engine; algorithm selection accordion; "
     "column picker; inline ResultsPanel; Run Model and Open Deep Analysis actions."),
    ("frontend/src/components/Dashboard/DeepAnalysisPage.jsx",
     "Full-screen deep research dashboard; Recharts visualisations; React-Three-Fiber 3D plot; "
     "AI Analyst Chat integration; URL parameter configuration."),
    ("frontend/src/utils/apiClient.js",
     "Centralised REST/WebSocket client; /api/ml/analyze and /api/ml/suggest call sites."),
    ("backend/app/services/db_connector.py",
     "Database abstraction layer; get_connection; async query executor."),
    ("backend/app/services/schema_analyzer.py",
     "Schema metadata cache; provides column types and row counts to /api/ml/suggest."),
]
add_table(
    ["File / Module", "Responsibility"],
    component_map,
    col_widths=[Inches(2.8), Inches(4.0)],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — FUTURE ROADMAP
# ══════════════════════════════════════════════════════════════════════════════

add_heading("12. Future Roadmap", level=1)
divider()

roadmap = [
    ("Version 2.0",
     "Predictive 3D Simulation",
     "Work on Data outcomes will be projected back onto the main Valkyrie 3D graph. "
     "Nodes will change size and position based on ML multi-day forecasts, allowing users "
     "to visualise 'Predicted Future States' of their database topology."),
    ("Version 2.1",
     "Cross-table Feature Joins",
     "secondary_tables parameter will be fully utilised to perform automatic SQL JOINs "
     "on detected foreign-key relationships, enriching the feature matrix with relational context."),
    ("Version 2.2",
     "Model Export & Scheduling",
     "Trained models (scikit-learn pickle) and analysis summaries (JSON) will be exportable. "
     "Scheduled re-runs will allow automated daily/weekly model refresh."),
    ("Version 2.3",
     "AutoML Hyperparameter Optimisation",
     "Bayesian optimisation (Optuna) will be integrated to automatically tune hyperparameters "
     "within a configurable time budget, replacing fixed defaults."),
    ("Version 3.0",
     "LLM-Native Query Interface",
     "Users will be able to describe their analytical goal in plain English ("
     "e.g., 'find the customers most likely to churn next month'), and the platform will "
     "automatically select the algorithm family, columns, and model."),
]
add_table(
    ["Version", "Feature", "Description"],
    roadmap,
    col_widths=[Inches(0.9), Inches(1.7), Inches(4.2)],
)

doc.add_paragraph()
divider()

# Footer paragraph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Living Data Intelligence Platform  —  Work on Data Documentation  —  Version 1.1  —  March 2026"
)
set_font(run, size=9, italic=True, color=(150, 150, 150))

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\karth\living-data-intelligence-backend-sasir\WORK_ON_DATA_COMPLETE_DOCUMENTATION.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
